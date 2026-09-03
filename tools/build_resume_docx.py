from __future__ import annotations

import argparse
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(119, 119, 119)
INK = RGBColor(34, 34, 34)
FONT_BODY = "DengXian"
FONT_HEAD = "Microsoft YaHei"
SECTIONS = {
    "个人简介",
    "核心能力",
    "项目与科研经历",
    "教育背景",
    "专业经验",
    "专业经验与荣誉",
    "获奖与荣誉",
    "技能",
    "参考资料",
}
ENTRY_SECTIONS = {"项目与科研经历", "教育背景", "专业经验"}
DATE_RE = re.compile(r"^\d{4}(?:[.\-]\d{2})?(?:[.\-]\d{2})?.*$")


def set_run_font(run, name=FONT_BODY, size=9.5, color=INK, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_table_fixed(table, widths_mm):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_mm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 56.6929)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_mm):
            cell.width = Mm(width)
            set_cell_margins(cell)


def set_bottom_border(paragraph, color="1F4E79", size="8", space="2"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def normalize_line(line):
    line = line.replace("\u00ad", "")
    line = line.replace("‐", "-").replace("‑", "-").replace("–", "-").replace("—", "-")
    line = line.replace("∙", "·")
    return re.sub(r"\s+", " ", line).strip()


def strip_marker(line):
    return re.sub(r"^[•◦▪●\-]\s*", "", line).strip()


def extract_lines(pdf_path, layout=False):
    command = ["pdftotext"]
    if layout:
        command.append("-layout")
    command.extend([str(pdf_path), "-"])
    raw = subprocess.check_output(command, text=True, encoding="utf-8", errors="replace")
    raw = raw.replace("\x0c", "\n")
    return [normalize_line(line) for line in raw.splitlines()]


def _load_identity() -> dict:
    """身份信息唯一事实源是 profile.json(个人数据,已 gitignore);缺省回退到示例档案。"""
    import json as _json
    from pathlib import Path as _Path
    base = _Path(__file__).resolve().parents[1] / "automation" / "profile"
    for name in ("profile.json", "profile.example.json"):
        try:
            data = _json.loads((base / name).read_text(encoding="utf-8"))
            return data.get("identity") or {}
        except (OSError, ValueError):
            continue
    return {}


IDENTITY = _load_identity()


def parse_resume(pdf_path):
    lines = extract_lines(pdf_path, layout=True)
    plain_lines = extract_lines(pdf_path, layout=False)
    first_section = next((i for i, line in enumerate(lines) if line in SECTIONS), len(lines))
    first_section_plain = next((i for i, line in enumerate(plain_lines) if line in SECTIONS), len(plain_lines))
    pre = [line for line in plain_lines[:first_section_plain] if line]
    location = (IDENTITY.get("location") or "").split("/")[0].strip()
    personal = [v for v in (IDENTITY.get("name"), IDENTITY.get("phone"), IDENTITY.get("email"), location) if v]
    header_noise = set(personal) | {"|"}
    summary_lines = [line for line in pre if line not in header_noise and not any(v in line for v in personal)]
    summary = "".join(summary_lines).strip()

    sections = []
    current = None
    for line in lines[first_section:]:
        if not line:
            continue
        if line in SECTIONS:
            current = {"name": line, "lines": []}
            sections.append(current)
        elif current is not None:
            current["lines"].append(line)

    return summary, sections


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(15.5)
    section.right_margin = Mm(15.5)
    section.top_margin = Mm(12.5)
    section.bottom_margin = Mm(12.5)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for style_name in ("List Bullet", "List Bullet 2"):
        style = doc.styles[style_name]
        style.font.name = FONT_BODY
        style.font.size = Pt(9.5)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(1)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.left_indent = Mm(5.5)
        style.paragraph_format.first_line_indent = Mm(-3.5)


def add_header(doc, photo_path):
    table = doc.add_table(rows=1, cols=2)
    remove_table_borders(table)
    set_table_fixed(table, [163, 25])
    left, right = table.rows[0].cells
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(IDENTITY.get("name") or "[姓名]")
    set_run_font(run, FONT_HEAD, 25, BLUE, True)

    p = left.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("   |   ".join(v for v in (IDENTITY.get("phone"), IDENTITY.get("email")) if v) or "[手机号]   |   [邮箱]")
    set_run_font(run, FONT_BODY, 9.5, GRAY)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(str(photo_path), width=Mm(22), height=Mm(29.7))

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after = Pt(5)
    set_bottom_border(rule, "1F4E79", "12", "1")


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, FONT_HEAD, 13.5, BLUE, True)
    set_bottom_border(p, "1F4E79", "8", "2")
    return p


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_bullet(doc, text, bold_label=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together = True
    if bold_label:
        match = re.match(r"(.+?[:：])\s*(.*)$", text)
        if match:
            r1 = p.add_run(match.group(1) + " ")
            set_run_font(r1, FONT_BODY, 9.5, INK, True)
            r2 = p.add_run(match.group(2))
            set_run_font(r2)
            return p
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_entry(doc, title, date, body_items):
    table = doc.add_table(rows=1, cols=2)
    remove_table_borders(table)
    set_table_fixed(table, [163, 25])
    set_cant_split(table.rows[0])
    left, right = table.rows[0].cells
    p = left.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    set_run_font(run, FONT_HEAD, 10.5, INK, True)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(date)
    set_run_font(run, FONT_BODY, 9, GRAY)

    for kind, text in body_items:
        if kind == "bullet":
            add_bullet(doc, text)
        else:
            add_body_paragraph(doc, text)


def coalesce_bullets(lines):
    result = []
    for line in lines:
        if line.startswith("• ") or line.startswith("- "):
            result.append(line)
        elif result:
            result[-1] += line
    return result


def add_section_content(doc, section):
    name = section["name"]
    lines = section["lines"]
    add_heading(doc, name)
    if name == "个人简介":
        add_body_paragraph(doc, "".join(lines))
        return
    if name not in ENTRY_SECTIONS:
        for line in coalesce_bullets(lines):
            if line.startswith("• "):
                add_bullet(doc, strip_marker(line), bold_label=(name == "核心能力"))
            elif line.startswith("- "):
                add_bullet(doc, strip_marker(line))
            else:
                add_body_paragraph(doc, line)
        return

    i = 0
    while i < len(lines):
        line = lines[i]
        looks_like_undecorated_entry = (
            i + 1 < len(lines)
            and DATE_RE.match(lines[i + 1])
            and not line.startswith(("- ", "• "))
        )
        if not line.startswith("• ") and not looks_like_undecorated_entry:
            if line.startswith("- "):
                add_bullet(doc, strip_marker(line))
            else:
                add_body_paragraph(doc, line)
            i += 1
            continue
        title = strip_marker(line)
        i += 1
        while i < len(lines) and not lines[i]:
            i += 1
        date = ""
        if i < len(lines) and DATE_RE.match(lines[i]):
            date = lines[i]
            i += 1
        body = []
        while i < len(lines):
            next_is_undecorated_entry = (
                i + 1 < len(lines)
                and DATE_RE.match(lines[i + 1])
                and not lines[i].startswith(("- ", "• "))
            )
            if lines[i].startswith("• ") or next_is_undecorated_entry:
                break
            part = lines[i]
            if part.startswith("- "):
                body.append(("bullet", strip_marker(part)))
            elif body and body[-1][0] == "bullet":
                body[-1] = ("bullet", body[-1][1] + part)
            elif part:
                body.append(("paragraph", part))
            i += 1
        add_entry(doc, title, date, body)


def build_docx(pdf_path, output_path, photo_path):
    summary, sections = parse_resume(pdf_path)
    doc = Document()
    setup_document(doc)
    add_header(doc, photo_path)
    if summary:
        add_body_paragraph(doc, summary)
    for section in sections:
        add_section_content(doc, section)

    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.title = f"{IDENTITY.get('name') or '[姓名]'} - 简历"
    props.subject = "求职简历"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-dir", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--photo", type=Path, required=True)
    args = parser.parse_args()

    pdfs = sorted(args.input_dir.glob("main_*.pdf"))
    if not pdfs:
        parser.error("--input-dir 中没有 main_*.pdf 可转换")
    for pdf in pdfs:
        out = args.output_dir / (pdf.stem + ".docx")
        build_docx(pdf, out, args.photo)
        print(out)


if __name__ == "__main__":
    main()
